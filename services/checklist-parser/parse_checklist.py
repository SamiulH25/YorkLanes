#!/usr/bin/env python3
"""
Parse York University degree checklist PDF or DOCX files.

Outputs JSON to stdout:
{
  "programme_hint": str | null,
  "years": [{ "year": 1, "courses": [{ "code", "credits", "raw", "section"? }] }],
  "warnings": [str]
}
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

# AP/ADMS 1000 3.00 | ADMS 3120 3.0 | LE/EECS 1011 3.00
COURSE_PATTERN = re.compile(
    r"(?<![A-Za-z])"
    r"(?:AP/|FA/|HH/|SC/|LE/|SB/|GL/|ES/)?"
    r"([A-Z]{2,6})\s+(\d{4})"
    r"(?:\s+([\d.]+))?",
    re.IGNORECASE,
)

YEAR_HEADER_PATTERN = re.compile(
    r"(?:"
    r"year\s*(four|[1-4])|"
    r"year\s*(\d)|"
    r"(\d)(?:st|nd|rd|th)\s+year|"
    r"(first|second|third|fourth)\s+year(?:\s+courses?)?"
    r")",
    re.IGNORECASE,
)

SECTION_YEAR_HINTS: list[tuple[re.Pattern[str], int]] = [
    (re.compile(r"general\s+education", re.I), 1),
    (re.compile(r"required\s+core|major\s*[–-]", re.I), 2),
    (re.compile(r"stream|specialization", re.I), 3),
    (re.compile(r"outside\s+the\s+major|credits\s+outside", re.I), 3),
    (re.compile(r"free\s+choice", re.I), 4),
]

# Section labels only — checklist year comes from explicit year headers, not these hints.
SECTION_LABEL_HINTS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"general\s+education", re.I), "General Education"),
    (re.compile(r"required\s+core|major\s*[–-]", re.I), "Required Core"),
    (re.compile(r"stream|specialization", re.I), "Stream / Specialization"),
    (re.compile(r"outside\s+the\s+major|credits\s+outside", re.I), "Outside the Major"),
    (re.compile(r"free\s+choice", re.I), "Free Choice"),
]

# Checklist sections that should become a single placeholder instead of importing every listed course.
# More specific patterns must come first.
STUB_SECTION_RULES: list[tuple[re.Pattern[str], str, str]] = [
    (re.compile(r"science\s+complementar|natural\s+science\s+complementar", re.I), "COMPLEMENTARY", "Science Complementary"),
    (re.compile(r"general\s+education|\bgen\s+ed\b", re.I), "GENERAL_ED", "General Education"),
    (re.compile(r"free\s+choice", re.I), "FREE_CHOICE", "Free Choice"),
    (re.compile(r"outside\s+the\s+major|credits\s+outside", re.I), "OUTSIDE_MAJOR", "Outside the Major"),
    (re.compile(r"breadth\s+requirement", re.I), "BREADTH", "Breadth Requirement"),
    (re.compile(r"humanities\s+requirement|social\s+science\s+requirement|natural\s+science\s+requirement", re.I), "BREADTH", "Breadth Requirement"),
    (re.compile(r"complementar", re.I), "COMPLEMENTARY", "Complementary Studies"),
    (re.compile(r"\belectives?\b", re.I), "ELECTIVE", "Electives"),
]

REQUIRED_SECTION_RULES: list[re.Pattern[str]] = [
    re.compile(r"required\s+core|core\s+courses?", re.I),
    re.compile(r"major\s*[–-]\s*required|programme\s+requirements?", re.I),
    re.compile(r"^\s*required\b", re.I),
]

LIST_REFERENCE_PATTERN = re.compile(
    r"(one of|choose from|select from|from the following|any of|pick\s+\d+|choose\s+\d+|"
    r"complete\s+\d+\s+credit|students?\s+must\s+select)",
    re.I,
)

CREDIT_REQUIREMENT_PATTERN = re.compile(
    r"(\d+(?:\.\d+)?)\s*(?:credit|credits|cr\.?)",
    re.I,
)

COURSE_CODE_PATTERN = re.compile(r"^[A-Z]{2,6} \d{4}$")

YEAR_WORDS = {"first": 1, "second": 2, "third": 3, "fourth": 4, "four": 4}

FULL_YEAR_PATTERN = re.compile(r"full[\s-]?year", re.I)

FOOTER_LINE_PATTERN = re.compile(
    r"(general\s+prerequisite|prerequisite:|grade\s+point|cross-listed|"
    r"note:\s*\"?major\"?|completed\s+major|common\)\s+prerequisite)",
    re.I,
)

BULLET_ONLY_PATTERN = re.compile(r"^[\uf0a8\u2022•\-\s]+$")

SUBJECT_BLOCKLIST = frozenset(
    {
        "THE",
        "FOR",
        "OR",
        "AND",
        "FROM",
        "WITH",
        "NOTE",
        "YEAR",
        "MUST",
        "TAKE",
        "LIST",
        "THAT",
        "THIS",
        "WHEN",
        "THAN",
        "INTO",
        "OVER",
        "ONLY",
        "ALSO",
        "THES",
        "MOST",
        "LEVEL",
    }
)

VALID_COURSE_NUMBER = re.compile(r"^[1-4]\d{3}$")

TOTAL_ROW_PATTERN = re.compile(r"^total\b", re.I)
CREDIT_CELL_PATTERN = re.compile(r"^\d+(?:\.\d+)?$")
SLOT_LINE_SUFFIX = re.compile(r"\|\s*(\d+(?:\.\d+)?)\s*$")
ANY_COURSE_PATTERN = re.compile(r"any\s+course", re.I)
OUTSIDE_CWR_PATTERN = re.compile(r"course\s+outside", re.I)
WORKSHOP_PATTERN = re.compile(r"workshop\s+credits?", re.I)
ADDITIONAL_MAJOR_PATTERN = re.compile(
    r"additional\s+credits?\s+chosen|creative\s+writing\s+list\s+of\s+courses",
    re.I,
)
GEN_ED_CATEGORY_PATTERN = re.compile(r"^\d+\)\s*", re.I)
PICK_LIST_PATTERN = re.compile(
    r"chosen from the following|choose from the following|select from the following",
    re.I,
)
MAJOR_SECTION_PATTERN = re.compile(r"major\s*[–-]\s*\d+\s*credits?", re.I)
TABLE_HEADER_PATTERN = re.compile(
    r"^(credit|complete|or incomplete|grade|notes)\b",
    re.I,
)

SKIP_LINE_PATTERN = re.compile(
    r"^(total|upper[\s-]?level|requirement|important|please|academic|department|registration|"
    r"note:)",
    re.IGNORECASE,
)


def normalize_code(subject: str, number: str) -> str:
    return f"{subject.upper()} {number}"


def parse_year_header(line: str) -> int | None:
    match = YEAR_HEADER_PATTERN.search(line)
    if not match:
        return None
    for group in match.groups():
        if not group:
            continue
        if group.isdigit():
            return int(group)
        word = group.lower()
        if word in YEAR_WORDS:
            return YEAR_WORDS[word]
    return None


def parse_section_year_hint(line: str) -> int | None:
    for pattern, year in SECTION_YEAR_HINTS:
        if pattern.search(line):
            return year
    return None


def parse_section_label_hint(line: str) -> str | None:
    for pattern, label in SECTION_LABEL_HINTS:
        if pattern.search(line):
            return label
    return None


def course_level(code: str) -> int:
    parts = code.split()
    if len(parts) < 2 or not parts[1][:1].isdigit():
        return 0
    return int(parts[1][0])


def should_treat_as_option_choice(code: str, checklist_year: int) -> bool:
    """Courses well below the checklist year level are usually pick-one options."""
    if checklist_year < 3:
        return False
    return course_level(code) < checklist_year


def is_likely_required_course_line(
    line: str,
    line_courses: list[dict],
    checklist_year: int,
) -> bool:
    """Single upper-level course rows are core requirements, not option lists."""
    if len(line_courses) != 1:
        return False
    code = line_courses[0]["code"]
    if course_level(code) >= max(1, checklist_year):
        return True
    if re.search(r"stream\s*[:\-–]", line, re.I) and course_level(code) >= 3:
        return True
    return False


def is_option_list_line(
    line: str,
    line_courses: list[dict],
    checklist_year: int,
) -> bool:
    if not line_courses:
        return False
    if is_likely_required_course_line(line, line_courses, checklist_year):
        return False
    if len(line_courses) >= 2:
        return True
    if LIST_REFERENCE_PATTERN.search(line):
        return True
    if len(line_courses) == 1 and should_treat_as_option_choice(line_courses[0]["code"], checklist_year):
        return True
    if len(line_courses) >= 3:
        return True
    return False


def apply_course_line_metadata(course: dict, line: str) -> dict:
    if FULL_YEAR_PATTERN.search(line):
        course["schedule_note"] = "full_year"
        course["section_label"] = course.get("section_label") or "Full year course"
    return course


def format_option_codes(codes: list[str], limit: int = 12) -> str:
    if not codes:
        return ""
    if len(codes) <= limit:
        return ", ".join(codes)
    shown = ", ".join(codes[:limit])
    return f"{shown}, +{len(codes) - limit} more"


def detect_stub_section(line: str) -> tuple[str, str] | None:
    for pattern, stub_type, label in STUB_SECTION_RULES:
        if pattern.search(line):
            return stub_type, label
    return None


def is_required_section(line: str) -> bool:
    return any(pattern.search(line) for pattern in REQUIRED_SECTION_RULES)


def parse_credit_requirement(line: str) -> float | None:
    match = CREDIT_REQUIREMENT_PATTERN.search(line)
    if not match:
        return None
    try:
        return float(match.group(1))
    except ValueError:
        return None


def is_concrete_course_code(code: str) -> bool:
    return bool(COURSE_CODE_PATTERN.match(code))


def make_stub_course(
    stub_type: str,
    label: str,
    section: str,
    credits: float | None,
    raw: str,
    option_codes: list[str] | None = None,
) -> dict:
    options = option_codes or []
    return {
        "code": stub_type,
        "credits": credits,
        "raw": raw[:200],
        "section": section,
        "kind": "stub",
        "stub_type": stub_type,
        "section_label": label,
        "option_codes": options,
        "title": format_option_codes(options) if options else None,
    }


def should_collapse_line_to_stub(line: str, course_count: int, in_stub_section: bool) -> bool:
    if PICK_LIST_PATTERN.search(line):
        return False
    if in_stub_section:
        return True
    if course_count >= 3 and LIST_REFERENCE_PATTERN.search(line):
        return True
    if course_count >= 5:
        return True
    return False


def parse_slot_suffix(line: str) -> tuple[str, float] | None:
    match = SLOT_LINE_SUFFIX.search(line)
    if not match:
        return None
    label = line[: match.start()].strip().rstrip("|").strip()
    try:
        credits = float(match.group(1))
    except ValueError:
        return None
    return label, credits


def classify_slot_stub(label: str) -> tuple[str, str]:
    if ANY_COURSE_PATTERN.search(label):
        return "FREE_CHOICE", "Free Choice"
    if OUTSIDE_CWR_PATTERN.search(label):
        return "OUTSIDE_MAJOR", "Outside the Major"
    if GEN_ED_CATEGORY_PATTERN.search(label) or re.search(
        r"humanities|natural science|social science",
        label,
        re.I,
    ):
        return "GENERAL_ED", "General Education"
    if ADDITIONAL_MAJOR_PATTERN.search(label):
        return "ADDITIONAL_MAJOR", "Major Electives"
    if WORKSHOP_PATTERN.search(label):
        return "WORKSHOP", "Workshop Credits"
    return "ELECTIVE", "Electives"


def default_entry_credits(entry: dict) -> float:
    credits = entry.get("credits")
    if credits is not None:
        return float(credits)
    if entry.get("kind") == "stub":
        return 3.0
    return 6.0


def assign_program_year(entry: dict, workshop_index: int = 0) -> int:
    stub_type = entry.get("stub_type", "")
    code = entry.get("code", "")
    raw = entry.get("raw", "")

    if stub_type == "GENERAL_ED" or code == "GENERAL_ED":
        return 1
    if entry.get("kind") != "stub" and course_level(code) <= 2:
        return 1
    if stub_type == "WORKSHOP" or WORKSHOP_PATTERN.search(raw):
        return 2 if workshop_index % 2 == 0 else 3
    if stub_type == "ADDITIONAL_MAJOR" or ADDITIONAL_MAJOR_PATTERN.search(raw):
        return 2 + (workshop_index % 2)
    if stub_type == "OUTSIDE_MAJOR" or OUTSIDE_CWR_PATTERN.search(raw):
        return 3
    if stub_type == "FREE_CHOICE":
        return 4
    return 2


def redistribute_credit_checklist(
    year_buckets: dict[int, list[dict]],
    saw_year_header: bool,
) -> dict[int, list[dict]]:
    if saw_year_header:
        return year_buckets

    flat_entries: list[dict] = []
    for year in sorted(year_buckets):
        flat_entries.extend(year_buckets[year])

    if len(flat_entries) < 6:
        return year_buckets

    total_credits = sum(default_entry_credits(entry) for entry in flat_entries)
    if total_credits < 90:
        return year_buckets

    redistributed: dict[int, list[dict]] = {1: [], 2: [], 3: [], 4: []}
    workshop_index = 0
    additional_index = 0

    for entry in flat_entries:
        raw = entry.get("raw", "")
        stub_type = entry.get("stub_type", "")
        if stub_type == "WORKSHOP" or WORKSHOP_PATTERN.search(raw):
            year = assign_program_year(entry, workshop_index)
            workshop_index += 1
        elif stub_type == "ADDITIONAL_MAJOR" or ADDITIONAL_MAJOR_PATTERN.search(raw):
            year = assign_program_year(entry, additional_index)
            additional_index += 1
        else:
            year = assign_program_year(entry)
        redistributed.setdefault(year, []).append(entry)

    return {year: courses for year, courses in redistributed.items() if courses}


def dedupe_merged_cells(cells: list[str]) -> list[str]:
    """Collapse Word horizontal merge repeats for the label column only."""
    cleaned = [cell.strip() for cell in cells if cell.strip()]
    if not cleaned:
        return []
    label = cleaned[0]
    rest = cleaned[1:]
    return [label, *rest]


def split_table_row(cells: list[str]) -> tuple[str, str]:
    cleaned = [cell.strip().replace("\n", " ") for cell in cells if cell.strip()]
    if not cleaned:
        return "", ""
    label = cleaned[0]
    credit_values: list[float] = []
    for cell in cleaned[1:]:
        credit_values.extend(parse_credit_values(cell))
    if credit_values:
        return label, " | ".join(f"{value:g}" for value in credit_values)
    if len(cleaned) > 1:
        return label, cleaned[1]
    return label, ""


def should_skip_duplicate_table_row(label: str, row_key: str, prev_row_key: str | None) -> bool:
    if prev_row_key is None or row_key != prev_row_key:
        return False
    if len(label) <= 60:
        return False
    if (
        PICK_LIST_PATTERN.search(label)
        or ADDITIONAL_MAJOR_PATTERN.search(label)
        or WORKSHOP_PATTERN.search(label)
    ):
        return True
    return False


def parse_credit_values(text: str) -> list[float]:
    values: list[float] = []
    for token in re.split(r"[|\s]+", text):
        token = token.strip()
        if CREDIT_CELL_PATTERN.match(token):
            values.append(float(token))
    return values


def normalize_table_row_key(label: str) -> str:
    return re.sub(r"\s+", " ", label.strip().lower())[:120]


def expand_label_credit_row(label: str, credit_text: str) -> list[str]:
    credits = parse_credit_values(credit_text)
    label = label.strip()
    if not label:
        return []

    if not credits:
        return [label]

    if len(credits) == 1 and COURSE_PATTERN.search(label):
        return [f"{label} | {credits[0]:g}"]

    return [f"{label} | {credit:g}" for credit in credits]


def extract_courses_from_line(line: str) -> list[dict]:
    found: list[dict] = []
    seen: set[str] = set()

    for match in COURSE_PATTERN.finditer(line):
        subject, number, credits_str = match.groups()
        subject = subject.upper()
        if subject in SUBJECT_BLOCKLIST or not VALID_COURSE_NUMBER.match(number):
            continue
        code = normalize_code(subject, number)
        if code in seen:
            continue
        seen.add(code)

        credits: float | None = None
        if credits_str:
            try:
                credits = float(credits_str)
            except ValueError:
                credits = None

        found.append(
            {
                "code": code,
                "credits": credits,
                "raw": match.group(0).strip(),
            }
        )

    return found


def merge_year_buckets(buckets: dict[int, list[dict]]) -> list[dict]:
    return [
        {"year": year, "courses": courses}
        for year, courses in sorted(buckets.items())
        if courses
    ]


def extract_courses_from_text(text: str) -> dict:
    warnings: list[str] = []
    current_year = 1
    current_section = "Requirements"
    year_buckets: dict[int, list[dict]] = {}
    seen_codes: set[str] = set()
    seen_stub_keys: set[str] = set()
    in_stub_section = False
    active_stub: dict | None = None
    saw_year_header = False

    def append_option_code(code: str, stub_type: str = "COMPLEMENTARY", label: str | None = None) -> None:
        nonlocal active_stub, in_stub_section
        if active_stub is None:
            active_stub = make_stub_course(
                stub_type,
                label or "Complementary Studies",
                current_section,
                parse_credit_requirement(current_section),
                current_section,
                [],
            )
            in_stub_section = True
        options = active_stub.setdefault("option_codes", [])
        if code not in options:
            options.append(code)
        active_stub["title"] = format_option_codes(options)

    def append_course(course: dict) -> None:
        code = course["code"]
        if course.get("kind") != "stub":
            if should_treat_as_option_choice(code, current_year):
                append_option_code(code)
                return
            if code in seen_codes:
                return
            seen_codes.add(code)
        else:
            stub_key = f"{current_year}:{code}:{len(year_buckets.get(current_year, []))}"
            if stub_key in seen_stub_keys:
                return
            seen_stub_keys.add(stub_key)
            # Merge consecutive complementary placeholders in the same year.
            year_list = year_buckets.setdefault(current_year, [])
            if (
                year_list
                and year_list[-1].get("kind") == "stub"
                and year_list[-1].get("code") == code
                and year_list[-1].get("stub_type") in {"COMPLEMENTARY", "ELECTIVE"}
                and not course.get("option_codes")
                and not year_list[-1].get("option_codes")
            ):
                existing = year_list[-1]
                for opt in course.get("option_codes", []):
                    opts = existing.setdefault("option_codes", [])
                    if opt not in opts:
                        opts.append(opt)
                if course.get("credits") is not None and existing.get("credits") is None:
                    existing["credits"] = course["credits"]
                existing["title"] = format_option_codes(existing.get("option_codes", []))
                return
        course.setdefault("kind", "course")
        course["section"] = current_section
        year_buckets.setdefault(current_year, []).append(course)

    def finalize_active_stub() -> None:
        nonlocal active_stub, in_stub_section
        if active_stub is not None:
            append_course(active_stub)
            active_stub = None
        in_stub_section = False

    def start_stub_section(line: str, stub_type: str, label: str) -> None:
        nonlocal active_stub, in_stub_section, current_section
        finalize_active_stub()
        current_section = line[:80]
        credits = parse_credit_requirement(line)
        active_stub = make_stub_course(stub_type, label, current_section, credits, line, [])
        in_stub_section = True

    def absorb_line_into_stub(line: str, line_courses: list[dict]) -> None:
        nonlocal active_stub
        if active_stub is None:
            return
        if active_stub.get("credits") is None:
            credits = parse_credit_requirement(line)
            if credits is not None:
                active_stub["credits"] = credits
        for course in line_courses:
            append_option_code(course["code"], active_stub["stub_type"], active_stub["section_label"])

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or BULLET_ONLY_PATTERN.fullmatch(line):
            continue
        if TABLE_HEADER_PATTERN.search(line):
            continue
        if re.search(r"^\d{4}\s+level", line, re.I):
            continue
        if SKIP_LINE_PATTERN.search(line) or FOOTER_LINE_PATTERN.search(line):
            continue

        year_num = parse_year_header(line)
        if year_num is not None and len(line) < 60 and not COURSE_PATTERN.search(line):
            finalize_active_stub()
            current_year = year_num
            saw_year_header = True
            continue

        stub_match = detect_stub_section(line)
        if stub_match is not None and COURSE_PATTERN.search(line) is None:
            stub_type, label = stub_match
            section_credits = parse_credit_requirement(line)
            if (
                stub_type in {"GENERAL_ED", "FREE_CHOICE", "OUTSIDE_MAJOR"}
                and section_credits is not None
                and section_credits >= 18
                and not PICK_LIST_PATTERN.search(line)
            ):
                finalize_active_stub()
                current_section = label
                continue
            start_stub_section(line, stub_type, label)
            continue

        if MAJOR_SECTION_PATTERN.search(line) and COURSE_PATTERN.search(line) is None:
            finalize_active_stub()
            current_section = "Required Core"
            continue

        if is_required_section(line) and COURSE_PATTERN.search(line) is None:
            finalize_active_stub()
            current_section = line[:80]
            label_hint = parse_section_label_hint(line)
            if label_hint:
                current_section = label_hint
            continue

        section_label = parse_section_label_hint(line)
        if section_label is not None and COURSE_PATTERN.search(line) is None:
            if detect_stub_section(line) is None:
                finalize_active_stub()
            current_section = section_label
            continue

        slot = parse_slot_suffix(line)
        if slot is not None:
            label, slot_credits = slot
            slot_courses = extract_courses_from_line(label)
            if slot_courses and PICK_LIST_PATTERN.search(label):
                finalize_active_stub()
                stub_type, stub_label = classify_slot_stub(label)
                append_course(
                    make_stub_course(
                        stub_type,
                        stub_label,
                        current_section,
                        slot_credits,
                        line,
                        [course["code"] for course in slot_courses],
                    )
                )
                continue
            if not slot_courses:
                finalize_active_stub()
                stub_type, stub_label = classify_slot_stub(label)
                append_course(
                    make_stub_course(
                        stub_type,
                        stub_label,
                        current_section,
                        slot_credits,
                        line,
                    )
                )
                continue
            if len(slot_courses) == 1 and is_concrete_course_code(slot_courses[0]["code"]):
                finalize_active_stub()
                course = slot_courses[0]
                course["credits"] = slot_credits
                apply_course_line_metadata(course, label)
                append_course(course)
                continue

        line_courses = extract_courses_from_line(line)

        if in_stub_section:
            if is_likely_required_course_line(line, line_courses, current_year):
                finalize_active_stub()
            elif is_option_list_line(line, line_courses, current_year):
                absorb_line_into_stub(line, line_courses)
                continue
            elif not line_courses:
                continue
            else:
                finalize_active_stub()

        if should_collapse_line_to_stub(line, len(line_courses), in_stub_section):
            stub_type = "ELECTIVE"
            label = "Electives"
            for pattern, candidate_type, candidate_label in STUB_SECTION_RULES:
                if pattern.search(current_section) or pattern.search(line):
                    stub_type = candidate_type
                    label = candidate_label
                    break
            credits = parse_credit_requirement(line) or parse_credit_requirement(current_section)
            option_codes = [course["code"] for course in line_courses]
            append_course(
                make_stub_course(
                    stub_type,
                    label,
                    current_section,
                    credits,
                    line,
                    option_codes,
                )
            )
            continue

        for course in line_courses:
            apply_course_line_metadata(course, line)
            append_course(course)

    finalize_active_stub()

    year_buckets = redistribute_credit_checklist(year_buckets, saw_year_header)
    years = merge_year_buckets(year_buckets)

    if not years or all(len(y["courses"]) == 0 for y in years):
        warnings.append(
            "No course codes found. Use a text-based PDF or DOCX checklist from your faculty site."
        )

    stub_count = sum(
        1 for year in years for course in year["courses"] if course.get("kind") == "stub"
    )
    if stub_count:
        warnings.append(
            f"Summarized {stub_count} complementary/elective checklist section(s) as placeholders "
            "(not every optional course was imported)."
        )

    programme_hint = None
    for line in text.splitlines()[:30]:
        stripped = line.strip()
        if not stripped:
            continue
        lower = stripped.lower()
        if any(
            token in lower
            for token in (
                "checklist",
                "honours",
                "bsc",
                "bcom",
                "beng",
                "bachelor",
                "commerce",
                "creative writing",
                "software engineering",
            )
        ):
            programme_hint = stripped[:200]
            break

    return {
        "programme_hint": programme_hint,
        "years": years if years else [{"year": 1, "courses": []}],
        "warnings": warnings,
    }


def read_pdf(path: Path) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        first_row_text = table.rows[0].cells[0].text.strip().lower() if table.rows else ""
        if "upper-level checklist" in first_row_text:
            continue

        prev_row_key: str | None = None
        for row in table.rows:
            raw_cells = [cell.text for cell in row.cells]
            cells = dedupe_merged_cells(raw_cells)
            label, credit_text = split_table_row(cells)
            if not label:
                continue
            if TABLE_HEADER_PATTERN.search(label) or label.lower().startswith("credit |"):
                continue
            if TOTAL_ROW_PATTERN.search(label):
                prev_row_key = None
                continue

            row_key = normalize_table_row_key(label)
            if should_skip_duplicate_table_row(label, row_key, prev_row_key):
                continue
            prev_row_key = row_key

            parts.extend(expand_label_credit_row(label, credit_text))

    return "\n".join(parts)


def parse_file(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = read_pdf(path)
    elif suffix in {".docx", ".doc"}:
        if suffix == ".doc":
            return {
                "programme_hint": None,
                "years": [{"year": 1, "courses": []}],
                "warnings": [".doc files are not supported. Save as .docx or export to PDF."],
            }
        text = read_docx(path)
    else:
        return {
            "programme_hint": None,
            "years": [{"year": 1, "courses": []}],
            "warnings": [f"Unsupported file type: {suffix}. Upload PDF or DOCX."],
        }

    return extract_courses_from_text(text)


def main() -> None:
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: parse_checklist.py <file>"}))
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(json.dumps({"error": f"File not found: {path}"}))
        sys.exit(1)

    try:
        result = parse_file(path)
        print(json.dumps(result))
    except Exception as exc:  # noqa: BLE001
        print(
            json.dumps(
                {
                    "error": str(exc),
                    "years": [{"year": 1, "courses": []}],
                    "warnings": [],
                }
            )
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
